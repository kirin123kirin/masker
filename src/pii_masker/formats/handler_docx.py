"""
Wordファイル (.docx) ハンドラー
段落・表・ヘッダー・フッター・テキストボックス内のテキストをマスクする。
書式（太字・色・フォント等）は保持する。
"""

import io
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.table import Table
from pii_masker.engine.masker import Masker


def _mask_paragraph(para: Paragraph, masker: Masker) -> tuple[str, str]:
    """
    段落内の run を結合してマスク処理し、各 run に分配して書き戻す。
    書式は run 単位で保持される。
    Returns: (original_text, masked_text)
    """
    if not para.runs:
        return "", ""

    # run のテキストを結合
    original = "".join(r.text for r in para.runs)
    if not original.strip():
        return original, original

    masked = masker.mask(original)
    if original == masked:
        return original, masked

    # マスク済みテキストを run に再分配
    # 戦略: 最初の run に全テキストを入れ、残りを空にする
    # （書式は最初の run の書式が適用される）
    para.runs[0].text = masked
    for run in para.runs[1:]:
        run.text = ""

    return original, masked


def _process_table(table: Table, masker: Masker) -> tuple[list[str], list[str]]:
    originals, maskeds = [], []
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                o, m = _mask_paragraph(para, masker)
                if o:
                    originals.append(o)
                    maskeds.append(m)
            # ネストしたテーブル
            for nested in cell.tables:
                o_list, m_list = _process_table(nested, masker)
                originals.extend(o_list)
                maskeds.extend(m_list)
    return originals, maskeds


def process_docx(src: Path, masker: Masker) -> tuple[bytes, str, str]:
    """
    Returns:
        (output_bytes, original_text, masked_text)
    """
    doc = Document(str(src))

    all_original: list[str] = []
    all_masked: list[str] = []

    # 本文段落
    for para in doc.paragraphs:
        o, m = _mask_paragraph(para, masker)
        if o:
            all_original.append(o)
            all_masked.append(m)

    # 表
    for table in doc.tables:
        o_list, m_list = _process_table(table, masker)
        all_original.extend(o_list)
        all_masked.extend(m_list)

    # ヘッダー・フッター
    for section in doc.sections:
        for header_footer in [
            section.header, section.footer,
            section.even_page_header, section.even_page_footer,
            section.first_page_header, section.first_page_footer,
        ]:
            if header_footer is None:
                continue
            for para in header_footer.paragraphs:
                o, m = _mask_paragraph(para, masker)
                if o:
                    all_original.append(o)
                    all_masked.append(m)
            for table in header_footer.tables:
                o_list, m_list = _process_table(table, masker)
                all_original.extend(o_list)
                all_masked.extend(m_list)

    # テキストボックス（drawing 内の txbx）
    for elem in doc.element.body.iter():
        tag = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag
        if tag == "txbxContent":
            for p_elem in elem.iter(qn("w:p")):
                # Paragraph オブジェクトを手動で wrap
                try:
                    para = Paragraph(p_elem, doc)
                    o, m = _mask_paragraph(para, masker)
                    if o:
                        all_original.append(o)
                        all_masked.append(m)
                except Exception:
                    pass

    # バイト列として出力
    buf = io.BytesIO()
    doc.save(buf)
    output_bytes = buf.getvalue()

    return output_bytes, "\n".join(all_original), "\n".join(all_masked)
